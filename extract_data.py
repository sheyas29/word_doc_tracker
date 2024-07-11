import docx
import json
import os
import sys
from tkinter import Tk
from tkinter.filedialog import askopenfilename

def extract_data_from_docx(docx_path):
    doc = docx.Document(docx_path)
    data = []
    id_counter = 1

    for table in doc.tables:
        for row in table.rows[1:]:  # Skip the header row
            is_bold = any(run.bold for cell in row.cells for paragraph in cell.paragraphs for run in paragraph.runs)
            si_no = row.cells[3].text.strip()
            confirmed_by = row.cells[5].text.strip()
            is_checked = bool(confirmed_by) or "." in si_no
            row_data = {
                'id': id_counter,
                'hrs': row.cells[0].text.strip(),
                'mins': row.cells[1].text.strip(),
                'sec': row.cells[2].text.strip(),
                'si.no': si_no,
                'activity_name': row.cells[4].text.strip(),
                'confirmed_by': confirmed_by,
                'is_bold': is_bold,
                'is_checked': is_checked
            }
            data.append(row_data)
            id_counter += 1

    return data

if __name__ == "__main__":
    Tk().withdraw()  # Prevents the Tk window from showing up
    docx_path = askopenfilename(title="Select the Word document", filetypes=[("Word Documents", "*.docx")])
    
    if not docx_path:
        print("No file selected. Exiting.")
        sys.exit()
    
    activities = extract_data_from_docx(docx_path)
    
    # Get the root directory of the project
    if getattr(sys, 'frozen', False):
        # If the script is run as an executable, get the path to the executable
        base_dir = os.path.dirname(sys.executable)
    else:
        # If the script is run from the source, get the path to the script
        base_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Define the root directory assuming the backend folder is at the root of the project
    root_dir = os.path.abspath(os.path.join(base_dir, '..', '..'))
    
    # Define the output path in the backend folder at the root of the project
    output_path = os.path.join(root_dir, 'backend', 'activities.json')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    with open(output_path, 'w') as f:
        json.dump(activities, f, indent=4)

    print(f"Data extracted and saved to {output_path}")
